"""Function to generate pdf"""
import io
import docx
from docx import Document
from docx.shared import Cm, Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def add_hyperlink(paragraph, text, url):
    """Add clickable email hyperlink"""
    if not text:
        paragraph.add_run("N/A")
        return

    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0000FF")
    r_pr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)

    new_run.append(r_pr)

    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)


def generate_docx(data):
    """
    Generates Form 'A' - Mediation Application Form DOCX
    :param data: dict with keys of form values
    :return: BytesIO buffer of DOCX
    """
    buffer = io.BytesIO()
    doc = Document()

    section = doc.sections[0]
    for attr, value in {
        "top_margin": 1,
        "bottom_margin": 1,
        "left_margin": 1.3,
        "right_margin": 1.3
    }.items():
        setattr(section, attr, Cm(value))

    title = doc.add_paragraph("FORM 'A'\nMEDIATION APPLICATION FORM\n[REFER RULE 3(1)]")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.bold = True
        run.font.size = Pt(13)

    auth = doc.add_paragraph("Mumbai District Legal Services Authority\nCity Civil Court, Mumbai")
    auth.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=16, cols=3)
    table.style = "Table Grid"
    widths = [0.5, 1.5, 5.6]
    for i, width in enumerate(widths):
        table.columns[i].width = Inches(width)

    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Cm(0.9)

        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)

    merged = table.rows[0].cells[0].merge(table.rows[0].cells[1]).merge(table.rows[0].cells[2])
    p = merged.paragraphs[0]
    run = p.add_run("DETAILS OF PARTIES:")
    run.bold = True
    run.font.size = Pt(12)

    table.cell(1, 0).text = "1"
    table.cell(1, 1).text = "Name of Applicant"
    table.cell(1, 2).text = data.get("client_name", "")

    row = table.rows[2].cells
    merged = row[1].merge(row[2])
    run = merged.paragraphs[0].add_run("Address and Contact Details of Applicant")
    run.bold = True
    run.font.size = Pt(11)

    registered = data.get("registered_address", data.get("branch_address", ""))
    correspondence = data.get("correspondence_address", data.get("branch_address", ""))

    table.cell(3, 1).text = "Address"
    addr_cell = table.cell(3, 2)
    addr_cell.text = ""
    for heading, value in [
        ("REGISTERED ADDRESS:\n", registered),
        ("CORRESPONDENCE ADDRESS:\n", correspondence)
    ]:
        p = addr_cell.add_paragraph()
        r = p.add_run(heading)
        r.bold = True
        r.font.size = Pt(11)
        p.add_run(value).font.size = Pt(10)
    addr_cell.add_paragraph()

    table.cell(4, 1).text = "Telephone No."
    table.cell(4, 2).text = data.get("tele_no", "")

    table.cell(5, 1).text = "Mobile No."
    table.cell(5, 2).text = data.get("mobile", "")

    table.cell(6, 1).text = "Email ID"
    email_p = table.cell(6, 2).paragraphs[0]
    email = data.get("email", "")
    add_hyperlink(email_p, email, f"mailto:{email}")

    row = table.rows[7].cells
    row[0].text = "2"
    merged = row[1].merge(row[2])
    run = merged.paragraphs[0].add_run("Name, Address & Contact Details of Opposite Party:")
    run.bold = True
    run.font.size = Pt(11)

    table.cell(8, 1).text = "Name"
    table.cell(8, 2).text = data.get("customer_name", "").upper()

    party_addr = data.get("customer_address", "").strip() or "________________________"
    table.cell(9, 1).text = "Address"
    cell = table.cell(9, 2)
    cell.text = ""
    for head in ["REGISTERED ADDRESS", "CORRESPONDENCE ADDRESS"]:
        p = cell.add_paragraph()
        r = p.add_run(f"{head}:\n")
        r.bold = True
        r.font.size = Pt(11)
        p.add_run(party_addr).font.size = Pt(10)
    cell.add_paragraph()

    table.cell(10, 1).text = "Telephone No."
    table.cell(10, 2).text = data.get("customer_tele_no", "")

    table.cell(11, 1).text = "Mobile No."
    table.cell(11, 2).text = data.get("customer_mobile", "")

    table.cell(12, 1).text = "Email ID"
    table.cell(12, 2).text = data.get("customer_email", "")

    dispute_head = table.rows[13].cells[0].merge(
        table.rows[13].cells[1]
    ).merge(table.rows[13].cells[2])
    rh = dispute_head.paragraphs[0].add_run("DETAILS OF DISPUTE:")
    rh.bold = True
    rh.font.size = Pt(12)

    d2 = table.rows[14].cells[0].merge(table.rows[14].cells[1]).merge(table.rows[14].cells[2])
    p2 = d2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("THE COMM. COURTS (PRE-INSTITUTION … SETTLEMENT) RULES, 2018")
    r2.bold = True

    final = table.rows[15].cells[1].merge(table.rows[15].cells[2])
    r3 = final.paragraphs[0].add_run(
        "Nature of dispute as per section 2(1)(c) of the Commercial Courts Act, 2015:"
    )
    r3.bold = True

    for i in range(1, 14):
        p = table.cell(i, 1).paragraphs[0]
        for r in p.runs:
            r.bold = True

    doc.save(buffer)
    buffer.seek(0)
    return buffer
